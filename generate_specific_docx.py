import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_specific_report():
    doc = Document()

    # Standard Academic 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette Definitions
    COLOR_PRIMARY = RGBColor(234, 88, 12)    # #EA580C (PrintHive Orange)
    COLOR_SECONDARY = RGBColor(15, 23, 42)   # #0F172A (Navy / Slate 900)
    COLOR_MUTED = RGBColor(100, 116, 139)    # #64748B (Slate 500)
    COLOR_DARK = RGBColor(30, 41, 59)        # #1E293B (Slate 800)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # =========================================================================
    # DOCUMENT HEADER / BANNER
    # =========================================================================
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    r_t = title_p.add_run("PRINTHIVE")
    r_t.font.name = "Calibri"
    r_t.font.size = Pt(28)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    r_s = sub_p.add_run("Decentralized 3-Sided Additive Manufacturing & 3D Commerce Grid\nCore Documentation: Abstract, Achievements & Report Organization")
    r_s.font.name = "Calibri"
    r_s.font.size = Pt(13)
    r_s.font.bold = True
    r_s.font.color.rgb = COLOR_SECONDARY

    # Divider Line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(20)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = RGBColor(226, 232, 240)

    # =========================================================================
    # 1. ABSTRACT (EXHAUSTIVE & DETAILED)
    # =========================================================================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("ABSTRACT")
    r_h1.font.name = "Calibri"
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_after = Pt(12)

    p_abs1 = doc.add_paragraph()
    p_abs1.paragraph_format.line_spacing = 1.25
    p_abs1.paragraph_format.space_after = Pt(12)
    p_abs1.add_run(
        "The rapid evolution of additive manufacturing (3D printing) technologies has unlocked immense potential for on-demand "
        "customized production, rapid physical prototyping, and distributed digital manufacturing. However, widespread consumer "
        "and enterprise adoption remains fundamentally constrained by severe structural fragmentation across the ecosystem. "
        "First, retail consumers and small enterprises seeking bespoke 3D items face prohibitive capital expenditures (₹15,000–₹80,000) "
        "for hardware acquisition, compounded by steep learning curves in Computer-Aided Design (CAD) modeling, slicing toolpaths, "
        "and thermal bed tuning. Second, freelance 3D designers lack a secure, automated monetization framework, frequently uploading "
        "high-value digital CAD assets to open web repositories without intellectual property protection or per-print royalties. "
        "Third, thousands of desktop and semi-industrial 3D printers owned by hobbyists, academic labs, and regional makerspaces "
        "suffer from severe hardware underutilization, remaining idle for 18 to 20 hours per day due to the absence of a reliable, "
        "localized order routing network."
    )

    p_abs2 = doc.add_paragraph()
    p_abs2.paragraph_format.line_spacing = 1.25
    p_abs2.paragraph_format.space_after = Pt(12)
    p_abs2.add_run(
        "To systematically address these interconnected bottlenecks, this project develops and deploys PrintHive — an AI-powered, "
        "decentralized, 3-sided hybrid commerce and on-demand manufacturing platform that bridges Buyers, 3D Designers, and Local 3D Printer Hubs "
        "under a collaborative and trustless operational model. PrintHive is engineered on a modern full-stack cloud architecture utilizing "
        "Next.js (App Router), TypeScript, PostgreSQL via Supabase with Row-Level Security (RLS), Three.js WebGL graphics, and Google Gemini AI."
    )

    p_abs3 = doc.add_paragraph()
    p_abs3.paragraph_format.line_spacing = 1.25
    p_abs3.paragraph_format.space_after = Pt(12)
    p_abs3.add_run(
        "The key technical innovations and subsystems comprising PrintHive include:\n"
        "1. In-Browser GPU-Accelerated 3D Viewport & Slicer: A Three.js/WebGL interactive rendering canvas that parses STL, 3MF, and OBJ files "
        "directly on the client device, calculating precise signed tetrahedron mesh volumes (cm³), raw polymer weight (grams), print durations, "
        "and automated manufacturing cost estimates across multiple filaments (PLA, PETG, ABS, TPU, Resin).\n"
        "2. Leaflet.js Geospatial Proximity Matching: A spatial matching engine that connects physical print orders to the nearest verified 3D printer "
        "farms within a 5–20 km radius using Haversine geodesic algorithms, collapsing delivery turnarounds from weeks to same-day or next-day local fulfillment.\n"
        "3. Razorpay Escrow State Machine & Tripartite Split: A trustless financial settlement pipeline that locks buyer funds in Escrow during "
        "manufacturing and releases an automated 70% / 15% / 15% revenue split (70% to Printer Hub, 15% to 3D Designer, 15% to Platform) upon buyer quality sign-off.\n"
        "4. AI-Driven Intelligence Engine: Google Gemini API integration enabling natural-language model searches and automated specification generation.\n"
        "5. Custom Brief Bidding & Dynamic Security Verification: A multi-tenant freelance CAD procurement portal featuring direct signed Cloudinary "
        "attachment pipelines and dynamic Cash-on-Delivery (COD) security verification tokens that eliminate bot-driven spam checkouts."
    )

    p_abs4 = doc.add_paragraph()
    p_abs4.paragraph_format.line_spacing = 1.25
    p_abs4.paragraph_format.space_after = Pt(24)
    p_abs4.add_run(
        "By transforming idle regional 3D printers into high-precision micro-factories and guaranteeing fair remuneration for creative IP, "
        "PrintHive establishes a sustainable, scalable, and decentralized additive manufacturing grid."
    )

    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_after = Pt(28)
    r_kw = p_kw.add_run("Keywords: Additive Manufacturing, Decentralized 3D Grid, Next.js, Three.js WebGL, Supabase RLS, Razorpay Escrow, Google Gemini AI, Leaflet Geolocation, Cloudinary CDN.")
    r_kw.font.bold = True
    r_kw.font.size = Pt(10)
    r_kw.font.color.rgb = COLOR_PRIMARY

    # =========================================================================
    # 2. SECTION 1.4 ACHIEVEMENTS (EXHAUSTIVE & SYSTEMATIC)
    # =========================================================================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("1.4 ACHIEVEMENTS & CORE TECHNICAL MILESTONES")
    r_h2.font.name = "Calibri"
    r_h2.font.size = Pt(18)
    r_h2.font.bold = True
    r_h2.font.color.rgb = COLOR_PRIMARY
    h2.paragraph_format.space_after = Pt(14)

    achievements_list = [
        ("1. Multi-Role 3-Sided Decentralized Architecture",
         "Architected and deployed a unified multi-tenant ecosystem supporting four distinct user roles: Buyers, 3D Designers (Creators), "
         "3D Printer Hub Owners, and Platform Administrators. Isolation between accounts, role-workspaces, and inventory is enforced strictly "
         "via server-side PostgreSQL Row Level Security (RLS) policies and authentication checks, utilizing client-side LocalStorage exclusively "
         "for active UI preferences and transient cart state."),
        
        ("2. Interactive In-Browser 3D Viewport & Automated Slicer Engine",
         "Engineered a high-performance WebGL 3D rendering pipeline utilizing Three.js and React Three Fiber (@react-three/fiber). "
         "The engine supports .stl, .3mf, .obj, and .gcode formats with orbit controls, interactive 3D perspective card physics, wireframe toggles, "
         "and real-time laser-scanning hologram effects. Implemented an automated mathematical slicer that calculates signed polygon volume (cm³), "
         "part mass in grams based on filament density constants (e.g., PLA = 1.24 g/cm³, PETG = 1.27 g/cm³, ABS = 1.04 g/cm³), and infill percentage "
         "to generate instant manufacturing cost estimates without server computational overhead."),

        ("3. Geospatial Proximity Matching & Leaflet.js Mapping Engine",
         "Implemented an interactive Leaflet.js and OpenStreetMap geolocation routing engine. Using spherical Haversine geodesic distance algorithms, "
         "the system dynamically calculates the physical proximity between buyer delivery coordinates and registered printer hub coordinates, "
         "automatically ranking and routing orders to verified 3D print farms within a 5–20 km radius to eliminate excessive transit delays and high logistics fees."),

        ("4. Escrow-Protected Payment State Machine & Fair 70/15/15 Revenue Split",
         "Integrated Razorpay Escrow payment gateway with automated cryptographic webhook verification. The platform enforces a tamper-proof "
         "tripartite economic split:\n"
         "  • 70% Direct Payout to Printer Hub: Compensates machine operational hours, electricity, polymer consumption, and packaging.\n"
         "  • 15% Royalty to 3D Designer: Compensates digital model IP creation, providing passive income whenever physical models are ordered.\n"
         "  • 15% Platform Maintenance Fee: Covers cloud infrastructure, AI inference, and dispute resolution.\n"
         "Funds remain securely held in Escrow until the buyer confirms delivery and quality compliance."),

        ("5. Google Gemini AI Generative Search & Intelligence Engine",
         "Integrated Google Gemini API (@google/genai) to power natural-language semantic searching (e.g., 'Find a durable heat-resistant drone mount in ABS'), "
         "automated technical description synthesis for uploaded models, and intelligent polymer/infill recommendations based on mechanical load constraints."),

        ("6. Two-Sided Custom 3D Design Brief & Bidding Exchange",
         "Built an end-to-end custom procurement engine. Buyers can publish custom CAD design briefs specifying dimensions, tolerance constraints, "
         "budget ranges, and multi-format attachments (.png, .jpg, .docx, .pdf, .stl). Verified freelance 3D designers and printer hubs review briefs "
         "and submit competitive bids (price and turnaround days), providing buyers with a centralized proposal review workspace."),

        ("7. Dynamic COD Security Token Gating & Hardened Upload Pipeline",
         "Developed an automated dynamic security verification system for Cash-on-Delivery (COD) checkouts that generates randomized 4-digit security tokens "
         "and locks the payment button until exact human verification is confirmed. In addition, built a direct signed Cloudinary media pipeline "
         "with universal support for images, CAD meshes, and technical documentation."),

        ("8. Luxury UI/UX Design System with Lucide React Vector Iconography",
         "Crafted a modern, responsive user interface following glassmorphic design principles with a warm pearl/cream palette (#FAF8F5), "
         "crisp slate typography (#0F172A), amber neon highlights (#FF6B35), and vector Lucide React outline icons housed in glowing squircle container badges.")
    ]

    for title, desc in achievements_list:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.line_spacing = 1.25
        p_item.paragraph_format.space_after = Pt(10)
        r_it = p_item.add_run(f"• {title}: ")
        r_it.font.bold = True
        r_it.font.color.rgb = COLOR_SECONDARY
        p_item.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 3. SECTION 1.5 ORGANIZATION OF REPORT (DETAILED ROADMAP)
    # =========================================================================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("1.5 ORGANIZATION OF REPORT")
    r_h3.font.name = "Calibri"
    r_h3.font.size = Pt(18)
    r_h3.font.bold = True
    r_h3.font.color.rgb = COLOR_PRIMARY
    h3.paragraph_format.space_after = Pt(14)

    p_org_intro = doc.add_paragraph()
    p_org_intro.paragraph_format.line_spacing = 1.25
    p_org_intro.paragraph_format.space_after = Pt(14)
    p_org_intro.add_run(
        "This technical project report is structured systematically across seven comprehensive chapters, progressing logically from "
        "theoretical motivation and market analysis to architectural design, mathematical modeling, implementation details, experimental validation, "
        "and future technological roadmaps. The organization is outlined as follows:"
    )

    chapters = [
        ("Chapter 1: Introduction",
         "Establishes the foundational context of additive manufacturing, highlights the tri-directional market failure (Buyers, Designers, Printer Hubs), "
         "defines the project objectives, outlines core technical achievements, and summarizes the structural organization of this documentation."),
        
        ("Chapter 2: Literature Survey & Market Analysis",
         "Presents a critical comparative review of existing centralized and repository-based 3D printing platforms (including Thingiverse, Shapeways, "
         "and Treatstock). Evaluates their limitations regarding localized fulfillment, intellectual property monetization, in-browser toolpath inspection, "
         "and escrow security, articulating the necessity for a collaborative 3-sided network."),

        ("Chapter 3: System Requirements & Feasibility Analysis",
         "Defines hardware and software prerequisites, technical feasibility, and functional requirements tailored to each of the four system roles "
         "(Buyer, Designer, Printer Hub, Administrator). Provides an in-depth justification for the selected technology stack (Next.js App Router, "
         "TypeScript, Three.js, PostgreSQL/Supabase, Razorpay Escrow, Google Gemini AI, and Cloudinary)."),

        ("Chapter 4: System Architecture & Database Design",
         "Outlines the distributed cloud architecture, component interaction pipelines, and the mathematical framework of the 70% / 15% / 15% escrow state machine. "
         "Provides Entity-Relationship Diagrams (ERDs), relational PostgreSQL table schema specifications with Row-Level Security (RLS) policies, "
         "and multi-level Data Flow Diagrams (DFD Level 0, Level 1, and Level 2)."),

        ("Chapter 5: Methodology & Implementation Details",
         "Details the concrete implementation of the core technical modules:\n"
         "  • Section 5.1: Three.js WebGL viewport rendering, STL/3MF binary parsing, and signed tetrahedron volume slicing algorithms.\n"
         "  • Section 5.2: Leaflet.js GPS geolocation proximity engine and Haversine routing algorithms.\n"
         "  • Section 5.3: Razorpay Escrow integration, webhook signature verification, and automated tripartite fund release.\n"
         "  • Section 5.4: Google Gemini AI natural-language prompt engineering and parameter extraction.\n"
         "  • Section 5.5: Custom 3D design brief bidding exchange and Cloudinary multi-format upload pipelines.\n"
         "  • Section 5.6: Dynamic Cash-on-Delivery (COD) security verification token gating and cart partition isolation."),

        ("Chapter 6: Testing, Results & Discussion",
         "Covers testing methodologies including TypeScript compile-time static type analysis, end-to-end user acceptance testing (UAT), "
         "client-side mesh parsing latency benchmarks across various polygon counts (10k to 500k triangles), security vulnerability validation, "
         "and RLS role boundary enforcement."),

        ("Chapter 7: Conclusion & Future Scope",
         "Summarizes the primary contributions and societal impact of the PrintHive ecosystem in democratizing additive manufacturing. "
         "Outlines future developmental trajectories, including IoT-connected OctoPrint/Klipper real-time printer telemetry, computer vision automated "
         "spaghetti/failure detection, blockchain-backed CAD IP watermarking, and cross-platform native mobile applications.")
    ]

    for ch_title, ch_desc in chapters:
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.line_spacing = 1.25
        p_ch.paragraph_format.space_after = Pt(12)
        r_cht = p_ch.add_run(f"• {ch_title}\n")
        r_cht.font.bold = True
        r_cht.font.size = Pt(11.5)
        r_cht.font.color.rgb = COLOR_DARK
        p_ch.add_run(ch_desc)

    # Save specifically to PrintHive_Abstract_Achievements_Organization.docx
    output_path = "PrintHive_Abstract_Achievements_Organization.docx"
    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    create_specific_report()
