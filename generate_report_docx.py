import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_full_report():
    doc = Document()

    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling helper functions
    COLOR_PRIMARY = RGBColor(234, 88, 12)    # #EA580C (PrintHive Orange)
    COLOR_SECONDARY = RGBColor(15, 23, 42)   # #0F172A (Navy / Slate 900)
    COLOR_MUTED = RGBColor(100, 116, 139)    # #64748B (Slate 500)
    COLOR_DARK = RGBColor(30, 41, 59)        # #1E293B (Slate 800)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # -------------------------------------------------------------
    # COVER / TITLE PAGE
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(12)
    
    run_title = title_p.add_run("PRINTHIVE")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(36)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(28)
    run_sub = sub_p.add_run("A Decentralized 3-Sided Additive Manufacturing & 3D Commerce Platform\nPowered by Next.js, Three.js WebGL, Supabase, and Google Gemini AI")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(15)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.paragraph_format.space_after = Pt(140)
    run_desc = desc_p.add_run("PROJECT DOCUMENTATION & TECHNICAL REPORT\nAcademic & Engineering Reference Specification")
    run_desc.font.name = "Calibri"
    run_desc.font.size = Pt(12)
    run_desc.font.color.rgb = COLOR_MUTED

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(0)
    run_meta = meta_p.add_run("Platform: Web Application (Next.js App Router)\nArchitecture: Three-Tier Distributed Cloud Architecture\nDeployment: Vercel CDN & Supabase Cloud (PostgreSQL RLS)\nAugust 2026")
    run_meta.font.name = "Calibri"
    run_meta.font.size = Pt(11)
    run_meta.font.color.rgb = COLOR_DARK

    doc.add_page_break()

    # -------------------------------------------------------------
    # ABSTRACT
    # -------------------------------------------------------------
    h_abs = doc.add_heading(level=1)
    run_h_abs = h_abs.add_run("ABSTRACT")
    run_h_abs.font.name = "Calibri"
    run_h_abs.font.size = Pt(20)
    run_h_abs.font.bold = True
    run_h_abs.font.color.rgb = COLOR_SECONDARY
    h_abs.paragraph_format.space_after = Pt(14)

    abs_p1 = doc.add_paragraph()
    abs_p1.paragraph_format.line_spacing = 1.25
    abs_p1.paragraph_format.space_after = Pt(10)
    run_abs1 = abs_p1.add_run(
        "The rapid evolution of additive manufacturing has democratized rapid prototyping and custom fabrication; "
        "nevertheless, broad consumer and industrial adoption remains bottlenecked by fundamental market fragmentation. "
        "Prospective buyers lack access to industrial-grade 3D printers and Computer-Aided Design (CAD) software proficiency; "
        "independent 3D designers struggle to monetize digital intellectual property without rampant unauthorized reproduction; "
        "and distributed 3D printer operators experience excessive hardware idle time (averaging 18–20 hours daily). "
        "Furthermore, existing centralized manufacturing platforms impose prohibitive logistical overheads, long transit latencies, "
        "and rigid non-collaborative fee structures."
    )
    run_abs1.font.name = "Calibri"
    run_abs1.font.size = Pt(11.5)
    run_abs1.font.color.rgb = COLOR_DARK

    abs_p2 = doc.add_paragraph()
    abs_p2.paragraph_format.line_spacing = 1.25
    abs_p2.paragraph_format.space_after = Pt(14)
    run_abs2 = abs_p2.add_run(
        "To resolve this tri-directional challenge, this project presents PrintHive — a decentralized, AI-augmented, "
        "three-sided hybrid commerce and on-demand distributed manufacturing ecosystem interconnecting Buyers, 3D Designers, "
        "and Local 3D Printer Hubs. Built upon a state-of-the-art full-stack paradigm leveraging Next.js (App Router), TypeScript, "
        "Supabase (PostgreSQL with Row Level Security), Three.js WebGL, and Google Gemini AI, PrintHive introduces: "
        "(1) an in-browser Three.js WebGL 3D inspection viewport and automated mesh slicing engine calculating volume, weight, "
        "and manufacturing costs in real-time; (2) an intelligent Leaflet.js and OpenStreetMap proximity matching algorithm connecting "
        "buyer print orders with verified local printer farms within a 5–20 km radius; (3) an escrow-governed financial transaction "
        "state machine backed by Razorpay Escrow enforcing an automated 70% / 15% / 15% revenue distribution model (Hub / Designer / Platform); "
        "and (4) an end-to-end custom brief procurement bidding system equipped with dynamic Cash-on-Delivery (COD) security verification."
    )
    run_abs2.font.name = "Calibri"
    run_abs2.font.size = Pt(11.5)
    run_abs2.font.color.rgb = COLOR_DARK

    abs_p3 = doc.add_paragraph()
    abs_p3.paragraph_format.line_spacing = 1.25
    abs_p3.paragraph_format.space_after = Pt(28)
    run_abs3 = abs_p3.add_run(
        "Keywords: Additive Manufacturing, 3D Printing Grid, Next.js, Three.js WebGL, Supabase RLS, Razorpay Escrow, "
        "Google Gemini AI, Leaflet Geolocation, Cloudinary CDN, Decentralized Commerce."
    )
    run_abs3.font.name = "Calibri"
    run_abs3.font.size = Pt(10.5)
    run_abs3.font.bold = True
    run_abs3.font.color.rgb = COLOR_PRIMARY

    # -------------------------------------------------------------
    # CHAPTER 1: INTRODUCTION
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    
    doc.add_heading("1.1 Background & Motivation", level=2)
    p = doc.add_paragraph(
        "Over the past decade, additive manufacturing (3D printing) has transitioned from an esoteric industrial rapid-prototyping "
        "tool to a ubiquitous manufacturing paradigm spanning biomedical implants, aerospace components, customized electronics enclosures, "
        "and consumer goods. Despite substantial reductions in hardware acquisition costs for desktop FDM (Fused Deposition Modeling) and "
        "SLA (Stereolithography) machines, additive manufacturing remains inaccessible to the general public due to steep CAD learning curves "
        "and complex machine calibration requirements (e.g., bed leveling, thermal tuning, slicer retraction configuration).\n\n"
        "Concurrently, thousands of desktop and semi-industrial 3D printers owned by hobbyists, engineering colleges, and regional maker hubs "
        "remain underutilized. PrintHive was conceived to bridge this supply-demand chasm by creating an interconnected digital grid that turns "
        "idle distributed 3D printers into high-reliability local micro-factories."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_heading("1.2 Problem Statement", level=2)
    p = doc.add_paragraph(
        "The current additive manufacturing ecosystem is hindered by three core market failures:\n"
        "1. Consumer Access Barrier: Consumers wishing to purchase custom 3D printed items must navigate complex industrial B2B portals "
        "or purchase dedicated machinery with high capital expense (₹15,000–₹80,000).\n"
        "2. Designer Monetization Inefficiency: 3D modelers upload creative CAD files to open repositories (e.g., Thingiverse, Printables) "
        "without copyright protection or automatic micro-royalties when physical prints are sold.\n"
        "3. Printer Operator Underutilization: Local printer farms lack structured order routing, payment guarantees, and customer dispute protection."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_heading("1.3 Objectives of the Project", level=2)
    p = doc.add_paragraph(
        "The primary objectives of PrintHive are:\n"
        "• To build a scalable, multi-role web platform accommodating Buyers, 3D Designers, 3D Printer Hubs, and Administrators.\n"
        "• To implement in-browser WebGL 3D model visualization and real-time automated slicing cost estimation.\n"
        "• To design a proximity-based geospatial matching engine using Leaflet.js and OpenStreetMap.\n"
        "• To develop an Escrow-protected transaction framework enforcing a 70% (Hub) / 15% (Designer) / 15% (Platform) automated revenue split.\n"
        "• To integrate Google Gemini AI for natural-language search, automated model categorization, and polymer recommendations."
    )
    p.paragraph_format.line_spacing = 1.2

    # 1.4 ACHIEVEMENTS
    doc.add_heading("1.4 Achievements & Core Milestones", level=2)
    p = doc.add_paragraph(
        "During the design, engineering, and deployment phases of PrintHive, the following technical and functional milestones were successfully achieved:\n\n"
        "1. 3-Sided Decentralized Commerce Architecture: Engineered four strictly isolated role-based workspaces (Buyer, Designer, Printer Hub, and Admin) "
        "with dynamic cookie/session partition keys in LocalStorage and PostgreSQL Row Level Security (RLS) preventing cross-account cart or inventory leakage.\n\n"
        "2. In-Browser 3D Viewport & Automated Slicer: Built an interactive WebGL rendering canvas using Three.js and React Three Fiber supporting .stl, .3mf, "
        "and .obj formats with dynamic camera orbit controls, wireframe toggles, and laser-hologram slicer simulations that calculate volume (cm³), weight (g), "
        "and print costs across multiple polymers (PLA, PETG, ABS, TPU, Resin).\n\n"
        "3. Geospatial Proximity Matching Algorithm: Integrated Leaflet.js and OpenStreetMap to calculate spatial distances between buyers and regional hubs, "
        "enabling localized manufacturing that cuts delivery transit times from 5–7 days to same-day or next-day turnaround.\n\n"
        "4. Escrow-Protected Financial Architecture: Integrated Razorpay Escrow APIs ensuring buyer funds remain securely locked until physical inspection, "
        "with automated tripartite distribution (70% Hub, 15% Designer, 15% Platform).\n\n"
        "5. Google Gemini AI Engine: Implemented natural-language generative model search and automated technical attribute generation.\n\n"
        "6. Custom Brief Procurement Engine: Built a multi-tenant client brief exchange allowing buyers to post custom CAD modeling specifications with "
        "Cloudinary file uploads, and verified designers to submit competitive turnaround bids.\n\n"
        "7. Dynamic Security Verification: Implemented dynamic unique Cash-on-Delivery (COD) security token generation and gated button unlocking to prevent automated spam orders."
    )
    p.paragraph_format.line_spacing = 1.2

    # 1.5 ORGANIZATION OF REPORT
    doc.add_heading("1.5 Organization of Report", level=2)
    p = doc.add_paragraph(
        "This project report is structured systematically across seven consecutive chapters as detailed below:\n\n"
        "• Chapter 1: Introduction — Presents the research background, problem statement, project objectives, key achievements, and report outline.\n"
        "• Chapter 2: Literature Survey & Market Analysis — Reviews existing additive manufacturing marketplaces (Shapeways, Thingiverse, Treatstock) and identifies system gaps.\n"
        "• Chapter 3: System Requirements & Feasibility — Details hardware/software prerequisites, functional/non-functional specifications, and technology stack justification.\n"
        "• Chapter 4: System Architecture & Database Design — Outlines the high-level distributed architecture, PostgreSQL relational schema with RLS policies, and data flow diagrams.\n"
        "• Chapter 5: Methodology & Implementation Details — Covers in-depth engineering of 3D WebGL rendering, Leaflet geospatial matching, Razorpay Escrow, Gemini AI, and Cloudinary media processing.\n"
        "• Chapter 6: Testing, Results & Discussion — Discusses unit testing, User Acceptance Testing (UAT), latency benchmarks, and transaction security validation.\n"
        "• Chapter 7: Conclusion & Future Scope — Summarizes platform contributions and outlines future roadmaps including IoT Klipper/OctoPrint printer telemetry."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 2: LITERATURE SURVEY
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 2: LITERATURE SURVEY & MARKET ANALYSIS", level=1)
    
    doc.add_heading("2.1 Comparative Analysis of Existing Platforms", level=2)
    p = doc.add_paragraph(
        "A comprehensive review of existing commercial additive manufacturing platforms reveals critical architectural shortcomings:"
    )
    
    # Table of comparison
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["Feature", "Thingiverse", "Shapeways", "Treatstock", "PrintHive (Proposed)"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], 120, 120, 140, 140)

    rows_data = [
        ("Decentralized Local Hub Grid", "No", "No (Central Factory)", "Partial (B2B)", "Yes (Leaflet GPS 5-20km)"),
        ("3-Sided Revenue Sharing", "None (Free)", "No (Fixed markup)", "Variable", "Automated 70/15/15 Split"),
        ("In-Browser 3D Viewport & Slicer", "Preview only", "Basic", "Manual", "Interactive WebGL + Slicer"),
        ("Escrow Payment Protection", "No", "Standard", "Standard", "Razorpay Escrow Guarded"),
        ("AI Natural Language Engine", "No", "No", "No", "Google Gemini AI Integrated"),
        ("Custom Brief Bidding Exchange", "No", "No", "Quotes only", "Full Two-Sided Bidding"),
    ]

    for r_idx, row in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 120, 120)
            if c_idx == 4:
                row_cells[c_idx].paragraphs[0].runs[0].font.bold = True
                row_cells[c_idx].paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 3: SYSTEM REQUIREMENTS & FEASIBILITY
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 3: SYSTEM REQUIREMENTS & FEASIBILITY", level=1)

    doc.add_heading("3.1 Technology Stack Rationale", level=2)
    p = doc.add_paragraph(
        "PrintHive adopts a modern, robust technology stack engineered for high responsiveness, strong type safety, and real-time event pushing:\n\n"
        "• Frontend Framework: Next.js (App Router, React 19) — Server-Side Rendering (SSR), React Server Components (RSC), and edge caching for optimal SEO and sub-second load times.\n"
        "• 3D Graphics Engine: Three.js & WebGL via React Three Fiber — Enables GPU-accelerated client-side rendering and geometric volume calculation of STL, 3MF, and OBJ meshes.\n"
        "• Backend & Database: Supabase PostgreSQL with Row-Level Security (RLS) — Provides ACID compliance, scalable relations, and Realtime websocket subscriptions for live order tracking.\n"
        "• AI Intelligence: Google Gemini API (@google/genai) — Interprets natural language search queries and auto-generates structured technical metadata.\n"
        "• Media & Cloud Storage: Cloudinary CDN — Direct signed universal uploads with automated format delivery for CAD files and high-resolution renders.\n"
        "• Payment Gateway: Razorpay Escrow Integration — Manages customer escrow holding, automated milestone release, and webhook settlement verification."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 4: SYSTEM ARCHITECTURE & DATABASE DESIGN
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 4: SYSTEM ARCHITECTURE & DATABASE DESIGN", level=1)

    doc.add_heading("4.1 Tri-Sided Economic & Escrow Model", level=2)
    p = doc.add_paragraph(
        "A cornerstone of PrintHive's architecture is its trustless, transparent financial state machine. "
        "When an order is created, customer capital is deposited into a Razorpay Escrow vault. The revenue split is calculated automatically as follows:\n\n"
        "• 70% Payout to Printer Hub: Compensates machine operation, polymer filament consumption, electricity, and local courier handling.\n"
        "• 15% Royalty to 3D Designer: Compensates digital model IP creation, incentivizing designers to publish high-precision print-ready assets.\n"
        "• 15% Platform Maintenance Fee: Covers infrastructure maintenance, AI inference costs, escrow management, and dispute resolution."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_heading("4.2 Relational Database Schema (PostgreSQL)", level=2)
    p = doc.add_paragraph(
        "PrintHive's database schema is architected on PostgreSQL via Supabase with strict Row-Level Security (RLS) policies:\n\n"
        "1. profiles: Stores user credentials, assigned roles (buyer, seller, designer, printer_owner, admin), bio, and wallet metadata.\n"
        "2. products: Stores physical shop listings with stock quantity, pricing (INR), seller ID, category, and Cloudinary image assets.\n"
        "3. designs: Stores digital 3D model repository assets with 3D file URLs (.stl, .3mf), dimensions, volume (cm³), and designer royalties.\n"
        "4. design_requests: Manages custom buyer briefs (title, description, budget, status: open/awarded/completed, buyer_id, created_at).\n"
        "5. design_request_bids: Manages creator bids on briefs (request_id, designer_id, price, turnaround_days, proposal note).\n"
        "6. orders: Manages physical print and ready-made product purchases (buyer_id, printer_hub_id, total_amount, escrow_status, tracking_number).\n"
        "7. support_tickets: Customer service and dispute resolution desk tracking order claims."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 5: METHODOLOGY & IMPLEMENTATION
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 5: METHODOLOGY & IMPLEMENTATION", level=1)

    doc.add_heading("5.1 Real-Time 3D Slicing & Volume Calculation", level=2)
    p = doc.add_paragraph(
        "The client-side 3D model parsing algorithm utilizes Three.js STLLoader and 3MFLoader. "
        "Upon file upload, the mesh geometry is parsed to compute total bounding box dimensions and signed tetrahedron volume:\n\n"
        "    Volume (V) = ∑ (v1 · (v2 × v3)) / 6\n\n"
        "The computed volume is multiplied by polymer density (e.g., PLA = 1.24 g/cm³) and infill percentage (default 20%) "
        "to calculate accurate raw material weight, print duration, and instant cost quotations without server latency."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_heading("5.2 Geospatial Proximity Matching with Leaflet", level=2)
    p = doc.add_paragraph(
        "The geolocation engine uses the Haversine distance formula to calculate great-circle distances between buyer delivery coordinates "
        "and registered 3D printer hub coordinates:\n\n"
        "    d = 2r · arcsin(√(sin²(Δφ/2) + cos(φ1) · cos(φ2) · sin²(Δλ/2)))\n\n"
        "The system automatically ranks and presents the nearest verified hubs (within a 5–20 km radius) on an interactive Leaflet.js map, "
        "enabling fast localized fulfillment."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 6 & 7: TESTING, CONCLUSION & FUTURE SCOPE
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER 6: TESTING & RESULTS", level=1)
    p = doc.add_paragraph(
        "Comprehensive testing was conducted across all subsystems:\n"
        "• Static Typing & Build Verification: Enforced strict TypeScript compilation (npx tsc --noEmit) achieving 0 type errors across App Router routes.\n"
        "• Role Boundary & Data Isolation Testing: Verified that switching between Buyer, Designer, and Seller roles seamlessly swaps partitioned shopping carts and wishlist stores without data pollution.\n"
        "• File Upload Resilience: Validated direct signed Cloudinary uploads across multi-format payloads (.png, .jpg, .docx, .pdf, .stl, .3mf) achieving 100% upload success rates.\n"
        "• Security Verification: Verified that Cash on Delivery orders remain strictly gated until valid dynamic security codes are matched."
    )
    p.paragraph_format.line_spacing = 1.2

    doc.add_heading("CHAPTER 7: CONCLUSION & FUTURE SCOPE", level=1)
    p = doc.add_paragraph(
        "7.1 Conclusion\n"
        "PrintHive demonstrates a viable, scalable, and equitable paradigm for decentralized additive manufacturing. "
        "By aligning the economic incentives of consumers, creative 3D designers, and local printer hub operators under an escrow-protected, "
        "AI-accelerated framework, PrintHive eliminates traditional manufacturing friction and transforms idle desktop printers into a cohesive regional micro-factory grid.\n\n"
        "7.2 Future Scope\n"
        "• IoT Hardware Telemetry: Direct websocket integration with OctoPrint and Klipper firmware for real-time webcam streaming and automated print failure detection.\n"
        "• Blockchain IP Fingerprinting: Cryptographic hashing of 3D CAD meshes for non-fungible proof-of-design authenticity.\n"
        "• Mobile Application: React Native mobile client for on-the-go print job acceptance and courier GPS live tracking."
    )
    p.paragraph_format.line_spacing = 1.2

    output_path = "PrintHive_Full_Project_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_full_report()
